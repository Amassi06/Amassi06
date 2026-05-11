from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path("rapport_activite")
OUT_DIR.mkdir(exist_ok=True)


CV_MD = """# Massinissa AIDEL

Développeur full-stack en alternance - Bachelor Architecture Logicielle

Choisy-le-Roi / Paris  
06 36 00 66 45  
aidelmassi@gmail.com  
Portfolio : https://aimassi.dev  
GitHub : https://github.com/Amassi06

## Profil

Développeur full-stack en Bachelor Architecture Logicielle à l'ESGI Paris, avec une expérience en alternance sur une application web complète pour SAS MOUSTAK / Matisse Food. Je conçois des interfaces React, des API REST Node.js/Express, des modèles de données PostgreSQL/MySQL et des déploiements Docker. Je recherche une alternance en développement logiciel/full-stack afin de renforcer mes compétences en architecture applicative, qualité de code et mise en production.

## Compétences techniques

Frontend : React, TypeScript, JavaScript, HTML, CSS, Tailwind CSS, Redux Toolkit, Ant Design, Vite
Backend : Node.js, Express, PHP, Laravel, API REST, JWT, bcrypt, validation de données
Bases de données : PostgreSQL, MySQL, SQL Server, Prisma, TypeORM, SQL
DevOps et outils : Docker, Docker Compose, Nginx, Git, GitHub, Swagger/OpenAPI, Postman, Figma
Méthodes et architecture : POO, MVC, SOLID, design patterns, clean code, Scrum, gestion par tickets, Notion, Trello
Autres technologies : Java, Spring Boot, Python, C#, C, n8n

## Expériences professionnelles

### Développeur full-stack en alternance - SAS MOUSTAK / Matisse Food
2026 - Aujourd'hui | Ivry-sur-Seine

- Conception et développement d'une application web de fidélisation client accessible sur https://matissefood.aimassi.dev.
- Développement d'un frontend React/Vite mobile-first : parcours client, validation de code, avis Google, roue de récompenses et page de résultat.
- Création d'une API REST Node.js/Express avec Prisma et PostgreSQL : codes uniques, lots, probabilités, statuts et paramètres publics.
- Développement d'une interface d'administration : authentification JWT, génération de codes, gestion des lots, statistiques et validation des récompenses en caisse.
- Mise en place du déploiement Docker/Docker Compose avec frontend Nginx, backend Node.js et base PostgreSQL.
- Suivi des tâches par tickets avec Notion, Trello et GitHub.

### Développeur full-stack - Stage JoyAtWork
Stage de 2 mois

- Développement de fonctionnalités web full-stack avec PHP, Laravel et Tailwind CSS.
- Participation à la conception d'interfaces, à l'intégration responsive et à la logique backend.
- Travail en environnement projet avec Git, structuration MVC et bonnes pratiques de développement.

## Projets

### Matisse Food - Application de fidélisation et roue de récompenses
https://matissefood.aimassi.dev

- Application full-stack React, Node.js, Express, Prisma, PostgreSQL et Docker pour un restaurant de restauration rapide.
- Parcours client avec code unique, lien d'avis Google, roue de tirage et affichage du lot gagné.
- Back-office sécurisé avec génération de codes, gestion des probabilités, statistiques et validation des lots.

### Zenith Cinéma - Plateforme cinéma full-stack
Application : https://zenith.aimassi.dev | API : https://api.aimassi.dev | Code : https://github.com/JEstermann/Cinema-NodeJS

- API REST Node.js/TypeScript avec Express, TypeORM, MySQL, JWT et documentation Swagger/OpenAPI.
- Fonctionnalités : authentification, CRUD films/salles/séances, wallet, achat de billets simples/super billets et historique d'utilisation.
- Frontend React/Vite avec interface de connexion, gestion des séances, billets, wallet et tableau de bord.
- Déploiement Docker Compose avec API, frontend Nginx, MySQL et configuration de sous-domaines.

### Assistant automatisé de gestion des emails et du calendrier

- Automatisation n8n connectée à OpenAI API, Gmail API, Google Calendar API et PostgreSQL.
- Analyse des emails entrants pour extraire réunions, deadlines, tâches et contacts.
- Création automatique d'événements Google Calendar avec rappels intégrés.

### Application de prévision météo - NASA Space Apps

- Application React, Node.js, Express, Leaflet et Chart.js consommant une API NASA.
- Visualisation de données météo, cartes interactives, graphiques dynamiques et export CSV/PDF.

## Formation

### ESGI Paris - Bachelor Architecture Logicielle
2025 - 2026 | 3ème année en alternance

- Formation orientée développement logiciel, architecture applicative, React, Node.js, bases de données, Docker, tests et méthodes agiles.
- Rythme : 3 semaines en entreprise / 1 semaine à l'école.

### Université Le Havre Normandie - Informatique
2023 - 2025 | Deux premières années d'études supérieures en informatique

- Bases en algorithmique, programmation, bases de données, développement web et conception logicielle.

### Baccalauréat série mathématique
Mention assez bien

## Langues

Français : courant
Anglais : avancé
Espagnol : débutant
"""


def set_cell_shading(paragraph, color="F2F4F7"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)
    set_cell_shading(p, "EEF3F8")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(1)
    p.add_run(text)


def add_heading_line(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(9.5)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.7)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    r = title.add_run("Massinissa AIDEL")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    rs = subtitle.add_run("Développeur full-stack en alternance - Bachelor Architecture Logicielle")
    rs.bold = True
    rs.font.size = Pt(10.5)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(5)
    contact.add_run(
        "Choisy-le-Roi / Paris | 06 36 00 66 45 | aidelmassi@gmail.com | "
        "https://aimassi.dev | https://github.com/Amassi06"
    )

    add_section(doc, "Profil")
    doc.add_paragraph(
        "Développeur full-stack en Bachelor Architecture Logicielle à l'ESGI Paris, avec une expérience en alternance "
        "sur une application web complète pour SAS MOUSTAK / Matisse Food. Je conçois des interfaces React, des API REST "
        "Node.js/Express, des modèles de données PostgreSQL/MySQL et des déploiements Docker. Je recherche une alternance "
        "en développement logiciel/full-stack afin de renforcer mes compétences en architecture applicative, qualité de code "
        "et mise en production."
    )

    add_section(doc, "Compétences techniques")
    skills = [
        "Frontend : React, TypeScript, JavaScript, HTML, CSS, Tailwind CSS, Redux Toolkit, Ant Design, Vite",
        "Backend : Node.js, Express, PHP, Laravel, API REST, JWT, bcrypt, validation de données",
        "Bases de données : PostgreSQL, MySQL, SQL Server, Prisma, TypeORM, SQL",
        "DevOps et outils : Docker, Docker Compose, Nginx, Git, GitHub, Swagger/OpenAPI, Postman, Figma",
        "Méthodes et architecture : POO, MVC, SOLID, design patterns, clean code, Scrum, gestion par tickets, Notion, Trello",
        "Autres technologies : Java, Spring Boot, Python, C#, C, n8n",
    ]
    for item in skills:
        add_bullet(doc, item)

    add_section(doc, "Expériences professionnelles")
    add_heading_line(doc, "Développeur full-stack en alternance - SAS MOUSTAK / Matisse Food", "2026 - Aujourd'hui | Ivry-sur-Seine")
    for item in [
        "Conception et développement d'une application web de fidélisation client accessible sur https://matissefood.aimassi.dev.",
        "Développement d'un frontend React/Vite mobile-first : parcours client, validation de code, avis Google, roue de récompenses et page de résultat.",
        "Création d'une API REST Node.js/Express avec Prisma et PostgreSQL : codes uniques, lots, probabilités, statuts et paramètres publics.",
        "Interface d'administration : authentification JWT, génération de codes, gestion des lots, statistiques et validation des récompenses en caisse.",
        "Déploiement Docker/Docker Compose avec frontend Nginx, backend Node.js et base PostgreSQL ; suivi par tickets Notion, Trello et GitHub.",
    ]:
        add_bullet(doc, item)

    add_heading_line(doc, "Développeur full-stack - Stage JoyAtWork", "Stage de 2 mois")
    for item in [
        "Développement de fonctionnalités web full-stack avec PHP, Laravel et Tailwind CSS.",
        "Participation à la conception d'interfaces, à l'intégration responsive et à la logique backend.",
        "Travail avec Git, structuration MVC et bonnes pratiques de développement.",
    ]:
        add_bullet(doc, item)

    add_section(doc, "Projets")
    add_heading_line(doc, "Matisse Food - Application de fidélisation et roue de récompenses", "https://matissefood.aimassi.dev")
    for item in [
        "Application full-stack React, Node.js, Express, Prisma, PostgreSQL et Docker pour un restaurant.",
        "Parcours client avec code unique, lien d'avis Google, roue de tirage et affichage du lot gagné.",
        "Back-office sécurisé avec génération de codes, gestion des probabilités, statistiques et validation des lots.",
    ]:
        add_bullet(doc, item)

    add_heading_line(doc, "Zenith Cinéma - Plateforme cinéma full-stack", "https://zenith.aimassi.dev | https://api.aimassi.dev | https://github.com/JEstermann/Cinema-NodeJS")
    for item in [
        "API REST Node.js/TypeScript avec Express, TypeORM, MySQL, JWT et documentation Swagger/OpenAPI.",
        "Fonctionnalités : authentification, CRUD films/salles/séances, wallet, achat de billets et historique d'utilisation.",
        "Frontend React/Vite et déploiement Docker Compose avec API, frontend Nginx et base MySQL.",
    ]:
        add_bullet(doc, item)

    add_heading_line(doc, "Assistant automatisé de gestion des emails et du calendrier", "n8n, OpenAI API, Gmail API, Google Calendar API, PostgreSQL")
    for item in [
        "Analyse des emails entrants pour extraire réunions, deadlines, tâches et contacts.",
        "Création automatique d'événements Google Calendar avec rappels intégrés.",
    ]:
        add_bullet(doc, item)

    add_section(doc, "Formation")
    add_heading_line(doc, "ESGI Paris - Bachelor Architecture Logicielle", "2025 - 2026 | 3ème année en alternance")
    add_bullet(doc, "Formation orientée développement logiciel, architecture applicative, React, Node.js, bases de données, Docker, tests et méthodes agiles.")
    add_bullet(doc, "Rythme : 3 semaines en entreprise / 1 semaine à l'école.")
    add_heading_line(doc, "Université Le Havre Normandie - Informatique", "2023 - 2025 | Deux premières années d'études supérieures en informatique")
    add_bullet(doc, "Bases en algorithmique, programmation, bases de données, développement web et conception logicielle.")
    add_heading_line(doc, "Baccalauréat série mathématique", "Mention assez bien")

    add_section(doc, "Langues")
    for item in ["Français : courant", "Anglais : avancé", "Espagnol : débutant"]:
        add_bullet(doc, item)

    path = OUT_DIR / "cv_aidel_massinissa_ats.docx"
    doc.save(path)
    return path


def write_markdown():
    path = OUT_DIR / "cv_aidel_massinissa_ats.md"
    path.write_text(CV_MD, encoding="utf-8")
    return path


if __name__ == "__main__":
    md = write_markdown()
    docx = build_docx()
    print(f"Generated: {md}")
    print(f"Generated: {docx}")
